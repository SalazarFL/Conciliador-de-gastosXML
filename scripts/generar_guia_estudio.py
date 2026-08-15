from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guia_de_estudio_Nexo_Fiscal.docx"

NAVY = "0C2461"
BLUE = "1E5AA8"
LIGHT_BLUE = "EAF1FB"
GOLD = "D4A017"
LIGHT_GOLD = "FFF6D8"
GREEN = "DFF3E4"
ORANGE = "FFF0D9"
RED = "FCE1E1"
GRAY = "F2F4F7"
DARK_GRAY = "3F4854"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Actualiza esta tabla en Word con clic derecho > Actualizar campo."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, end])


def add_run_text(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_rich_paragraph(doc, parts, style=None, align=None, space_after=5):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    for part in parts:
        if isinstance(part, str):
            add_run_text(p, part)
        else:
            text, attrs = part
            add_run_text(p, text, **attrs)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run_text(p, bold_prefix, bold=True, color=NAVY)
        add_run_text(p, text[len(bold_prefix):])
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_number(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE, border=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={"val": "single", "sz": 8, "color": border},
                    bottom={"val": "single", "sz": 8, "color": border},
                    left={"val": "single", "sz": 18, "color": border},
                    right={"val": "single", "sz": 8, "color": border})
    p = cell.paragraphs[0]
    add_run_text(p, title + ". ", bold=True, color=NAVY)
    add_run_text(p, text)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_flow(doc, steps):
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, step in enumerate(steps):
        cell = table.add_row().cells[0]
        set_cell_shading(cell, LIGHT_BLUE if i % 2 == 0 else LIGHT_GOLD)
        set_cell_border(cell, top={"val": "single", "sz": 5, "color": "B7C6DA"},
                        bottom={"val": "single", "sz": 5, "color": "B7C6DA"},
                        left={"val": "single", "sz": 5, "color": "B7C6DA"},
                        right={"val": "single", "sz": 5, "color": "B7C6DA"})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_text(p, step, bold=True if i in (0, len(steps) - 1) else False, color=NAVY)
        if i < len(steps) - 1:
            arrow_cell = table.add_row().cells[0]
            p2 = arrow_cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_text(p2, "↓", bold=True, color=GOLD, size=13)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_module_card(doc, purpose, inputs, process, output, storage):
    rows = [
        ("Para qué sirve", purpose),
        ("Qué recibe", inputs),
        ("Qué hace", process),
        ("Qué entrega", output),
        ("Dónde guarda", storage),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], NAVY)
        r = cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        cells[1].paragraphs[0].add_run(value)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 19, NAVY),
                              ("Heading 2", 14, BLUE), ("Heading 3", 11.5, GOLD)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(10.2)

    if "Termino" not in styles:
        term_style = styles.add_style("Termino", WD_STYLE_TYPE.CHARACTER)
        term_style.font.name = "Aptos"
        term_style.font.bold = True
        term_style.font.color.rgb = RGBColor.from_string(NAVY)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_text(header, "Nexo Fiscal  |  Guía de estudio", bold=True, color=NAVY, size=8.5)
    add_page_number(section.footer.paragraphs[0])

    # Pide a Word actualizar automáticamente el índice al abrir el archivo.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_document():
    doc = Document()
    configure_document(doc)

    # Portada
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "XML", bold=True, color=GOLD, size=24)
    add_run_text(p, "Concilia", bold=True, color=NAVY, size=24)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Guía de estudio del sistema")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(sub, "Cómo funciona cada módulo, explicado desde cero", color=BLUE, size=15)
    doc.add_paragraph()
    add_callout(doc, "Objetivo", "Entender el recorrido de la información: desde el correo o un archivo XML hasta la verificación semanal y los reportes, sin necesitar experiencia previa en programación.", LIGHT_GOLD, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "Preparada a partir del código real del proyecto", italic=True, color=DARK_GRAY, size=10)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p2, "Revisión: 20 de julio de 2026", color=DARK_GRAY, size=9)
    doc.add_page_break()

    doc.add_heading("Índice", level=1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    # 1
    doc.add_heading("1. La idea general del sistema", level=1)
    add_rich_paragraph(doc, [
        "Nexo Fiscal es una aplicación web que ayuda a responder una pregunta de negocio: ",
        ("“¿Las facturas que se planea pagar tienen un comprobante electrónico que las respalde y sus datos coinciden?”", {"bold": True, "color": NAVY}),
        " Para responderla, reúne información procedente del correo, de archivos XML y de listados CSV/XLSX, la organiza por empresa y semana, y compara los documentos."
    ])
    add_callout(doc, "La metáfora más útil", "Piensa en el sistema como una mesa de revisión. A un lado está el listado de lo que se desea pagar; al otro, las facturas electrónicas. Nexo Fiscal coloca cada factura junto a su línea probable, muestra las diferencias y deja los casos dudosos para decisión humana.")

    doc.add_heading("1.1 Flujo operativo actual", level=2)
    add_flow(doc, [
        "1. Iniciar sesión y elegir la sociedad activa",
        "2. Elegir o crear una semana de trabajo",
        "3. Subir el listado de facturas por pagar",
        "4. Obtener XML desde Correo o Carga de Facturas XML",
        "5. Importar y asignar los XML a la misma semana",
        "6. El sistema vuelve a verificar automáticamente el listado",
        "7. Revisar faltantes/diferencias y exportar resultados",
    ])
    add_callout(doc, "Regla central", "El listado de una semana se compara con las facturas XML asignadas a esa misma semana, no con todo el histórico. Esto reduce cruces equivocados.", LIGHT_GOLD, GOLD)

    doc.add_heading("1.2 Dos flujos que conviven", level=2)
    add_table(doc, ["Flujo", "Estado", "Módulos"], [
        ("Operativo actual", "Visible en el menú", "Inicio, Facturas por pagar, Facturas desde Correo, Carga de Facturas XML, Reportes y Usuarios"),
        ("Histórico o legado", "Sigue disponible por URL", "Carga de Gastos y Conciliación factura-vs-gasto"),
    ], [3.5, 3.5, 9.0])
    add_rich_paragraph(doc, ["Que un módulo sea ", ("legado", {"bold": True}), " no significa que esté borrado: sus rutas, controladores, modelos y vistas siguen en el proyecto. Simplemente no forma parte del menú principal actual."])

    # 2 concepts
    doc.add_heading("2. Conceptos importantes, en lenguaje sencillo", level=1)
    concepts = [
        ("Aplicación web", "Programa que se usa desde el navegador. El navegador solicita una página o una acción al servidor."),
        ("Servidor", "En este proyecto es principalmente Apache + PHP dentro de XAMPP. Recibe solicitudes y ejecuta la lógica."),
        ("PHP", "Lenguaje que ejecuta el sistema en el servidor."),
        ("MySQL / base de datos", "Lugar estructurado donde se conservan usuarios, facturas, semanas, listados, resultados y otros registros."),
        ("MVC", "Modelo-Vista-Controlador: una forma de separar datos, decisiones y pantalla."),
        ("Ruta", "Dirección interna, por ejemplo /correo o /por-pagar. Indica qué controlador debe atender una solicitud."),
        ("Controlador", "Coordina una acción: recibe datos, valida, llama modelos/helpers y decide qué respuesta entregar."),
        ("Modelo", "Clase que lee y escribe información en la base de datos."),
        ("Vista", "Plantilla PHP/HTML que produce lo que la persona ve en pantalla."),
        ("Helper o servicio", "Pieza reutilizable con una tarea especializada, como leer XML, comparar facturas o manejar la cola."),
        ("Sesión", "Memoria temporal asociada al usuario conectado; conserva identidad y semana activa."),
        ("GET", "Solicitud normalmente usada para abrir o consultar una página."),
        ("POST", "Solicitud usada normalmente para enviar datos o ejecutar un cambio."),
        ("AJAX / JSON", "Forma en que una pantalla habla con el servidor sin recargar toda la página. JSON es el formato de la respuesta."),
        ("IMAP", "Protocolo usado para leer buzones de correo desde el servidor."),
        ("XML", "Archivo estructurado de una factura electrónica. Contiene clave, emisor, receptor, fecha, montos y más."),
        ("Hash SHA-256", "Huella digital del contenido de un archivo; ayuda a detectar duplicados o contenido alterado."),
        ("Matching", "Proceso de buscar qué registro corresponde a otro usando número, proveedor, monto y otras señales."),
        ("Score", "Puntuación de similitud de 0 a 100. No es una certeza absoluta: resume qué tan parecidos son dos datos."),
        ("Cola de importación", "Lista de archivos pendientes de procesar por tandas para evitar que una carga grande agote el tiempo del servidor."),
    ]
    add_table(doc, ["Término", "Significado"], concepts, [4.0, 12.0])

    # 3 architecture
    doc.add_heading("3. Cómo viaja una solicitud por el sistema", level=1)
    add_flow(doc, [
        "Navegador: la persona pulsa un botón o abre una URL",
        "public/index.php: punto de entrada de la aplicación",
        "App.php: carga configuración, errores y rutas",
        "Router.php: encuentra la ruta GET/POST correspondiente",
        "Controlador: valida y coordina la operación",
        "Modelo / helper: consulta BD, lee archivos o aplica reglas",
        "Vista HTML o respuesta JSON: resultado que vuelve al navegador",
    ])
    doc.add_heading("3.1 Las tres piezas de MVC", level=2)
    add_table(doc, ["Pieza", "Pregunta que responde", "Ejemplo del proyecto"], [
        ("Modelo", "¿Cómo leo o guardo los datos?", "Factura.php consulta facturas_xml; PorPagar.php consulta listados y líneas."),
        ("Vista", "¿Cómo lo presento?", "views/correo/index.php construye la pantalla del módulo Correo."),
        ("Controlador", "¿Qué debe ocurrir ante esta acción?", "CorreoController.php valida, procesa adjuntos y manda datos a la vista o devuelve JSON."),
    ], [3.0, 5.2, 7.8])
    doc.add_heading("3.2 Piezas base", level=2)
    add_bullet(doc, "App.php inicia el sistema, detecta la URL base, configura la zona horaria y maneja errores globales.")
    add_bullet(doc, "Router.php compara la URL y el método HTTP con app/config/routes.php, y llama Controlador@método.")
    add_bullet(doc, "Controller.php ofrece renderizado, JSON, redirecciones, mensajes temporales, autenticación y semana activa.")
    add_bullet(doc, "Model.php crea una conexión PDO reutilizable y ejecuta consultas preparadas contra MySQL.")
    add_callout(doc, "Por qué las variables aparecen de repente en una vista", "Controller::render() usa extract($data). Las claves enviadas por el controlador se convierten en variables de la vista. Por ejemplo, 'cuentas' pasa a ser $cuentas. El editor puede marcarla como “indefinida” al analizar la vista aislada, aunque el flujo real sí la entrega.")

    # 4 cross-cutting
    doc.add_heading("4. Contextos compartidos que gobiernan el sistema", level=1)
    doc.add_heading("4.1 Usuario y sesión", level=2)
    add_rich_paragraph(doc, ["Casi todos los controladores ejecutan ", ("requireAuth()", {"bold": True, "color": NAVY}), ". Si no existe un user_id en la sesión, la persona es enviada a /login. La sesión también guarda nombre, usuario, rol de administrador y semana activa."])
    doc.add_heading("4.2 Sociedad activa", level=2)
    add_rich_paragraph(doc, ["Representa la empresa con la que se está trabajando. Solo una sociedad queda marcada como ", ("activa", {"bold": True}), ". Su cédula se usa en Correo para comprobar que el receptor del XML sea la empresa correcta. También se asocia al listado de pago."])
    doc.add_heading("4.3 Semana activa", level=2)
    add_rich_paragraph(doc, ["Es el contenedor lógico del trabajo semanal. Se guarda en la sesión bajo ", ("semana_activa", {"bold": True}), " y se comparte entre Carga XML, Correo y Facturas por pagar. Una factura sin semana queda fuera de una verificación que esté limitada a una semana específica."])
    doc.add_heading("4.4 Cuenta de correo activa", level=2)
    add_rich_paragraph(doc, ["El módulo Correo admite varias cuentas IMAP. La elegida se guarda en la configuración local y cada registro de índice/bandeja tiene ", ("cuenta_id", {"bold": True}), " para impedir que se mezclen buzones."])
    add_table(doc, ["Contexto", "Dónde se conserva", "Qué afecta"], [
        ("Usuario", "Sesión PHP + tabla usuarios", "Acceso, nombre visible y permisos de administrador"),
        ("Sociedad activa", "Tabla sociedades", "Cédula receptora y asociación del listado"),
        ("Semana activa", "Sesión + tabla semanas", "Filtro y universo del matching semanal"),
        ("Cuenta de correo", "correo_cuentas + config.json local", "Buzón que se busca/sincroniza/procesa"),
    ], [3.5, 5.0, 7.5])

    # 5 auth
    doc.add_heading("5. Módulo de autenticación", level=1)
    add_module_card(doc,
        "Controlar quién entra al sistema.",
        "Usuario, contraseña y token CSRF.",
        "Busca el usuario, verifica la contraseña cifrada con password_verify, comprueba que esté activo y crea la sesión.",
        "Acceso al Inicio o un mensaje de error.",
        "Tabla usuarios y sesión PHP.")
    doc.add_heading("5.1 Inicio de sesión paso a paso", level=2)
    for text in [
        "Al abrir /login, se genera un token CSRF aleatorio. Este token ayuda a comprobar que el formulario nació dentro de la aplicación.",
        "Al enviar el formulario, el servidor compara el token recibido con el de la sesión.",
        "Busca el username en la tabla usuarios.",
        "password_verify compara la contraseña escrita con el hash guardado; la contraseña original no necesita almacenarse.",
        "Si el usuario está activo, se regenera el identificador de sesión para reducir el riesgo de fijación de sesión.",
        "Se registra la fecha de último acceso y se redirige al Inicio.",
    ]:
        add_number(doc, text)
    add_callout(doc, "Cerrar sesión", "Vacía y destruye la sesión; luego redirige a /login.", GRAY, DARK_GRAY)

    # 6 home/society
    doc.add_heading("6. Inicio y Sociedades", level=1)
    add_module_card(doc,
        "Ser el panel de entrada y definir con qué empresa se trabaja.",
        "La sesión del usuario y las sociedades registradas.",
        "Cuenta facturas, pendientes de correo, obtiene el último listado por pagar y su resumen.",
        "Indicadores, sociedad activa y estado del último listado.",
        "sociedades, facturas_xml, correo_bandeja y tablas porpagar_*.")
    doc.add_heading("6.1 Lógica de sociedades", level=2)
    add_bullet(doc, "La primera sociedad creada se activa automáticamente.")
    add_bullet(doc, "Al activar una sociedad, el modelo primero desactiva la anterior y luego activa la elegida.")
    add_bullet(doc, "Nombre y cédula son obligatorios; la cédula debe contener números.")
    add_bullet(doc, "Si se elimina la sociedad activa y no queda otra activa, el sistema pide elegir una nueva.")
    add_callout(doc, "Importante", "Cambiar la sociedad activa cambia la cédula contra la que Correo valida las nuevas facturas procesadas. No reescribe automáticamente las facturas ya importadas.", LIGHT_GOLD, GOLD)

    # 7 weeks
    doc.add_heading("7. Semanas de trabajo", level=1)
    add_module_card(doc,
        "Separar el trabajo por períodos y evitar comparar contra todo el histórico.",
        "Una semana existente, una nueva semana o la opción sin semana.",
        "Resuelve/crea el registro y recuerda la selección en la sesión.",
        "Un semana_id que se asigna a facturas y listados.",
        "Tabla semanas; semana_id en facturas_xml y porpagar_listados.")
    add_rich_paragraph(doc, ["La semana no se calcula automáticamente a partir de la fecha de factura. Es una ", ("clasificación de trabajo elegida por la persona", {"bold": True, "color": NAVY}), ". Por eso una factura con fecha anterior puede pertenecer al lote de pago de la semana actual si así se asigna."])
    add_callout(doc, "Efecto automático", "Cuando una factura entra o sale de una semana, el sistema intenta volver a verificar los listados de ambas semanas. La asignación no falla aunque esa re-verificación auxiliar encuentre un problema.")

    # 8 por pagar
    doc.add_heading("8. Facturas por pagar — módulo central del flujo actual", level=1)
    add_module_card(doc,
        "Comprobar que cada línea del listado semanal de pagos tenga una factura XML de respaldo.",
        "Archivo CSV/XLSX, sociedad activa y semana elegida.",
        "Previsualiza, valida, elimina duplicados, guarda líneas y busca el mejor XML disponible de la misma semana.",
        "Checklist con estados respaldada, con diferencia o sin respaldo; también CSV exportable.",
        "porpagar_listados y porpagar_facturas; vínculos a facturas_xml.")

    doc.add_heading("8.1 Qué archivos acepta", level=2)
    add_bullet(doc, "Formato plano CSV/XLSX con columnas Fecha, Numero, Proveedor y Total. La fecha es informativa; número, proveedor y total válido son esenciales.")
    add_bullet(doc, "Reporte agrupado: reconoce encabezados del tipo “Proveedor <código> <nombre>” y líneas de documento tipo FACT-…; el sistema lo aplana automáticamente.")
    add_bullet(doc, "El formato .xls antiguo se rechaza; debe guardarse como .xlsx o .csv.")

    doc.add_heading("8.2 Vista previa e importación", level=2)
    for text in [
        "El archivo se guarda temporalmente y se analiza sin escribir líneas en la base de datos.",
        "Cada fila se clasifica como nueva, repetida o error.",
        "Al confirmar, el mismo análisis se reutiliza; así la vista previa y la importación real siguen las mismas reglas.",
        "Si la semana ya tiene listado, se agregan únicamente líneas nuevas al listado existente.",
        "Después de guardar, se ejecuta el matching contra las facturas XML de esa semana.",
    ]:
        add_number(doc, text)
    add_rich_paragraph(doc, [
        ("Detección de repetidas: ", {"bold": True, "color": NAVY}),
        "una línea se considera repetida si ya existe la combinación número + proveedor normalizados, o número + monto exacto. La segunda combinación cubre casos donde el nombre del proveedor viene recortado o escrito distinto."
    ])

    doc.add_heading("8.3 Cómo busca el XML que respalda una línea", level=2)
    add_flow(doc, [
        "Tomar una línea del listado aún sin vínculo manual",
        "Examinar facturas FE disponibles de la misma semana que todavía no se usaron",
        "Comparar número contra número corto y consecutivo completo",
        "Comparar el nombre del proveedor",
        "Elegir la candidata con mayor puntuación 60% número + 40% proveedor",
        "Comparar montos y asignar el estado",
    ])
    add_table(doc, ["Regla", "Valor actual", "Qué significa"], [
        ("Umbral de número", "90", "Solo se aceptan coincidencias fuertes del número."),
        ("Umbral de proveedor", "60", "El proveedor debe parecerse razonablemente."),
        ("Rescate", "Número ≥ 95 + monto dentro de ₡1", "Permite aceptar nombres comerciales/razones sociales muy diferentes cuando número y monto identifican la factura."),
        ("Peso final", "60% número + 40% proveedor", "El número tiene mayor importancia."),
        ("Tolerancia de monto", "₡1.00", "Diferencias de hasta un colón se consideran redondeo."),
    ], [4.0, 4.0, 8.0])
    add_callout(doc, "Una factura, una línea", "Durante el matching automático, una factura XML no puede respaldar dos líneas del mismo listado. Las facturas ya utilizadas se excluyen de las siguientes búsquedas.", LIGHT_GOLD, GOLD)

    doc.add_heading("8.4 Estados del checklist", level=2)
    add_table(doc, ["Estado", "Interpretación", "Acción recomendada"], [
        ("Respaldada", "Se encontró XML y el monto está dentro de la tolerancia.", "Revisión rápida; normalmente está lista."),
        ("Con diferencia", "Se encontró un XML probable, pero el total del listado y el XML difieren más de ₡1.", "Comprobar monto, notas de crédito o error del listado."),
        ("Sin respaldo", "No apareció una factura suficientemente compatible.", "Buscarla en Correo, cargar XML, mover de semana o vincular manualmente."),
    ], [3.2, 7.2, 5.6])

    doc.add_heading("8.5 Casos sin coincidencia y vínculo manual", level=2)
    add_rich_paragraph(doc, ["El botón de casos sin coincidencia reúne dos lados: las facturas XML de la semana que no respaldan ninguna línea y las líneas todavía sin respaldo. La persona puede escoger el par correcto. Ese vínculo queda marcado como ", ("match_manual", {"bold": True}), " para que una verificación automática posterior no lo deshaga."])
    add_callout(doc, "Qué sigue evaluándose al vincular manualmente", "Número y proveedor dejan de decidir porque la persona ya confirmó la relación. El sistema compara el monto para clasificarla como respaldada o con diferencia.")

    doc.add_heading("8.6 Buscar en Correo y exportar", level=2)
    add_bullet(doc, "Desde una línea se puede abrir Correo con un término derivado del número de factura y navegar por el listado con flechas.")
    add_bullet(doc, "La exportación genera un CSV separado por punto y coma, con BOM UTF-8 para que Excel respete acentos.")
    add_bullet(doc, "El archivo exportado contiene datos del listado, datos del XML vinculado, diferencia y estado.")

    # 9 correo
    doc.add_heading("9. Facturas desde Correo", level=1)
    add_module_card(doc,
        "Encontrar facturas en buzones IMAP, validar sus adjuntos y llevar los XML aprobados a la importación.",
        "Cuenta de correo, criterio de búsqueda, mensajes seleccionados, sociedad activa y semana.",
        "Sincroniza encabezados, descarga adjuntos, clasifica XML, verifica Hacienda/cédula, empareja PDF y encola XML.",
        "Bandeja de facturas pendientes/rechazadas; importación a Facturas XML y archivos FE_/NC_ renombrados.",
        "correo_cuentas, correo_indice, correo_carpetas, correo_procesados, correo_bandeja, storage/correo y cola de importación.")

    doc.add_heading("9.1 Configuración y cuentas IMAP", level=2)
    add_bullet(doc, "Se pueden registrar varias cuentas con host, puerto, usuario, contraseña y carpeta inicial.")
    add_bullet(doc, "La cuenta activa determina qué índice local y qué buzón se consultan.")
    add_bullet(doc, "La extensión IMAP de PHP debe estar activa; sin ella el módulo informa que no está disponible.")
    add_bullet(doc, "La carpeta destino local define dónde se copian los pares XML/PDF ya renombrados.")
    add_callout(doc, "Seguridad de la contraseña", "El modelo guarda la contraseña IMAP codificada en Base64, no cifrada. Base64 se puede revertir; la propia implementación asume una herramienta local. No debe tratarse como protección criptográfica.", RED, "C24141")

    doc.add_heading("9.2 Índice local y búsqueda", level=2)
    add_rich_paragraph(doc, ["Para que buscar sea rápido, el sistema mantiene en MySQL un ", ("índice local", {"bold": True, "color": NAVY}), " con encabezados: carpeta, UID, remitente, asunto y fecha. No descarga todos los adjuntos durante una búsqueda normal."])
    add_table(doc, ["Ámbito", "Fuente", "Comportamiento"], [
        ("Asunto / remitente / ambos", "Índice local MySQL", "Búsqueda rápida. Si el índice está vacío, se construye en el primer uso."),
        ("Todo, incluido cuerpo", "Servidor IMAP", "Más lenta porque consulta el contenido en el buzón carpeta por carpeta."),
        ("Mes prioritario", "Índice local", "Busca primero en el mes de la factura; si no hay resultados, amplía automáticamente."),
    ], [4.0, 4.0, 8.0])
    add_rich_paragraph(doc, ["La sincronización se hace por ", ("tandas", {"bold": True}), " de tiempo. Cada petición avanza alrededor de 20 segundos y la pantalla repite hasta terminar. Un lock global permite una sola sincronización a la vez, evitando conexiones IMAP duplicadas cuando hay varios usuarios o una tarea programada."])

    doc.add_heading("9.3 Procesamiento de correos seleccionados", level=2)
    add_flow(doc, [
        "Seleccionar correos (máximo 10 por petición del navegador)",
        "Descargar XML y PDF del mensaje",
        "Separar factura electrónica de MensajeHacienda",
        "Cruzar ambos por la Clave de 50 dígitos",
        "Validar aceptación/rechazo y cédula del receptor",
        "Emparejar cada PDF con su XML",
        "Guardar cada documento válido en la bandeja",
    ])
    add_rich_paragraph(doc, [
        ("MensajeHacienda: ", {"bold": True, "color": NAVY}),
        "no se importa como una factura adicional. Sirve para conocer la respuesta de Hacienda: código 1 aceptado, 2 aceptación parcial y 3 rechazado. Si un proveedor envía solo MensajeHacienda aceptado + PDF, el parser puede reconstruir los datos principales de la factura usando la Clave."
    ])

    doc.add_heading("9.4 Validaciones y emparejamiento del PDF", level=2)
    add_bullet(doc, "Cédula: si la sociedad activa tiene cédula y el receptor legible es diferente, el estado pasa a otra_cedula.")
    add_bullet(doc, "Duplicado: el hash del XML se compara contra facturas importadas y contra la propia bandeja.")
    add_bullet(doc, "PDF: si hay un XML y un PDF se emparejan directamente. Si hay varios, se comparan números extraídos del nombre, la Clave y la regla de terminación numérica.")
    add_bullet(doc, "PDF huérfano: si no puede asociarse a un XML, se mueve a storage/correo/sin_identificar y se reporta un aviso.")
    add_bullet(doc, "Factura rechazada u otra cédula: su PDF se descarta para que no llegue a las carpetas de trabajo.")

    doc.add_heading("9.5 Estados de la bandeja", level=2)
    add_table(doc, ["Estado", "Qué significa", "¿Se importa?"], [
        ("pendiente", "Factura nueva lista para revisión/importación.", "Sí"),
        ("importada", "Ya fue enviada a una importación XML.", "Ya ocurrió"),
        ("descartada", "La persona la retiró de la bandeja.", "No"),
        ("ya_existe", "Estado histórico; actualmente los duplicados se purgan o ni entran.", "No"),
        ("rechazada", "MensajeHacienda indicó rechazo.", "No"),
        ("otra_cedula", "El receptor no coincide con la sociedad activa.", "No"),
    ], [3.0, 9.0, 4.0])

    doc.add_heading("9.6 Importar desde la bandeja", level=2)
    for text in [
        "Se toman únicamente filas pendientes cuyo XML todavía existe en disco.",
        "Se resuelve la semana elegida y se crea una importación llamada “Correo <fecha/hora>”.",
        "Los XML se copian a la cola de importación; la bandeja queda marcada como importada.",
        "El navegador llama al procesador de cola por tandas hasta terminar.",
        "Si hay carpeta destino, se copian XML y PDF renombrados como FE_PROVEEDOR_ddmmaa_numero (o NC para nota de crédito).",
        "Si el par se copió correctamente, se borran los originales temporales. Si algo falla, se conservan para no perderlos.",
    ]:
        add_number(doc, text)
    add_callout(doc, "Descartar", "Al descartar, se borra el XML de trabajo, se conserva el PDF según la lógica actual y se elimina la marca de correo procesado para permitir procesar nuevamente el mensaje de origen.", GRAY, DARK_GRAY)

    # 10 invoice
    doc.add_heading("10. Carga de Facturas XML", level=1)
    add_module_card(doc,
        "Convertir archivos XML en registros de factura consultables y utilizables en el matching.",
        "Uno o varios XML y una semana elegida.",
        "Valida archivos, interpreta el XML, crea/recupera proveedor, evita duplicados y guarda datos fiscales.",
        "Facturas almacenadas, historial de importaciones y detalle consultable.",
        "importaciones, importacion_items, proveedores y facturas_xml.")

    doc.add_heading("10.1 Qué extrae el parser", level=2)
    add_table(doc, ["Dato", "Uso"], [
        ("Clave / consecutivo", "Identificación del comprobante y derivación del número corto."),
        ("Tipo de documento", "Distingue FE, NC y ND; por-pagar solo usa FE como respaldo."),
        ("Emisor", "Cédula/RFC y razón social para el catálogo de proveedores."),
        ("Receptor", "Permite comprobar la cédula de la empresa en el flujo de correo."),
        ("Fecha", "Información del documento y señal auxiliar en conciliación legado."),
        ("Subtotal, IVA y total", "Comparaciones y reportes."),
        ("Moneda", "Reporte y lectura financiera."),
        ("Contenido + hash", "Permite conservar el XML en BD, detectar duplicados y verificar integridad."),
    ], [5.0, 11.0])

    doc.add_heading("10.2 Importación directa y cola", level=2)
    add_table(doc, ["Modalidad", "Cuándo sirve", "Cómo trabaja"], [
        ("Directa", "Lotes pequeños", "PHP recibe todos los archivos y los procesa en una sola solicitud."),
        ("Cola", "Lotes grandes o importación desde Correo", "Crea items pendientes y procesa grupos pequeños repetidamente, mostrando progreso."),
    ], [3.5, 5.0, 7.5])
    add_rich_paragraph(doc, ["Los estados de un item de cola incluyen pendiente, procesando, importado, duplicado, sin_plantilla y error. Si un item queda en procesando demasiado tiempo, puede volver a pendiente para recuperarse de una ejecución interrumpida."])

    doc.add_heading("10.3 Prevención de duplicados", level=2)
    add_bullet(doc, "Primero compara el hash SHA-256 del documento.")
    add_bullet(doc, "También compara consecutivo, proveedor y fecha cuando corresponde.")
    add_bullet(doc, "Un duplicado se cuenta y reporta, pero no se inserta como una segunda factura.")
    add_callout(doc, "Después de importar", "Si entraron facturas nuevas en una semana, los listados por pagar de esa semana se re-verifican automáticamente.")

    # 11 reports
    doc.add_heading("11. Reportes y exportación", level=1)
    add_module_card(doc,
        "Consultar y descargar información filtrada.",
        "Tipo de reporte y filtros de fechas, estado, moneda, importación, corrida y columnas.",
        "Obtiene datos, aplica filtros, calcula un resumen y prepara filas para Excel.",
        "Vista previa JSON y archivo .xlsx.",
        "Lee facturas_xml, gastos_consolidados, conciliaciones e importaciones.")
    add_table(doc, ["Tipo", "Contenido", "Filtros destacados"], [
        ("Facturas", "Fecha, número, proveedor, subtotal, IVA, total, moneda y archivo.", "Fechas, moneda e importaciones."),
        ("Gastos", "Fecha, número, proveedor, items, base, IVA y total.", "Fechas e importaciones."),
        ("Conciliación legado", "Ambos lados, diferencias, score, tipo de match y hallazgos.", "Estado, match, fechas, corrida y solo diferencias."),
    ], [3.3, 7.4, 5.3])
    add_bullet(doc, "La vista previa devuelve como máximo 500 filas y avisa si fue truncada.")
    add_bullet(doc, "La exportación puede incluir solo las columnas seleccionadas.")
    add_bullet(doc, "En conciliación, la columna Hallazgos explica diferencias de proveedor, número, fecha o total.")
    add_bullet(doc, "El XLSX aplica estilos visuales a los hallazgos y ajusta el ancho de esa columna.")

    # 12 users
    doc.add_heading("12. Gestión de Usuarios", level=1)
    add_module_card(doc,
        "Crear, editar, activar/desactivar y eliminar cuentas de acceso.",
        "Nombre, username, email, contraseña, estado activo y rol administrador.",
        "Valida campos y duplicados; genera hash de contraseña en el modelo.",
        "Usuarios administrables y permisos de acceso.",
        "Tabla usuarios.")
    add_bullet(doc, "Solo un administrador puede entrar a este módulo; requireAdmin() comprueba el rol de la sesión.")
    add_bullet(doc, "Username y email deben ser únicos.")
    add_bullet(doc, "La contraseña debe tener al menos 6 caracteres y coincidir con su confirmación.")
    add_bullet(doc, "El usuario conectado no puede eliminarse a sí mismo, quitarse su propio rol de administrador ni desactivarse.")
    add_bullet(doc, "Al editar el propio nombre/username se actualizan también los valores mostrados en la sesión.")

    # 13 supporting catalog
    doc.add_heading("13. Catálogo de Proveedores", level=1)
    add_rich_paragraph(doc, ["No es una pantalla principal del menú; funciona como módulo de apoyo. Al importar una factura, ", ("Proveedor::obtenerOCrear", {"bold": True}), " busca al emisor por cédula/RFC y, si no existe, lo crea. Así muchas facturas pueden apuntar a un mismo proveedor."])
    add_bullet(doc, "Conserva razón social y una versión normalizada para búsquedas/matching.")
    add_bullet(doc, "La API /api/proveedores/buscar requiere al menos dos caracteres y devuelve coincidencias para autocompletado.")
    add_bullet(doc, "El catálogo evita repetir el nombre completo del emisor en cada fila de factura y facilita estadísticas.")

    # 14 legacy expenses
    doc.add_heading("14. Carga de Gastos — flujo legado", level=1)
    add_callout(doc, "Estado del módulo", "Sigue funcionando mediante /gastos, pero está fuera del menú actual porque el flujo principal fue reemplazado por Facturas por pagar.", ORANGE, GOLD)
    add_module_card(doc,
        "Importar registros contables de gastos para la conciliación histórica factura-vs-gasto.",
        "CSV o XLSX con Fecha, Numero, Proveedor, Iva y Total.",
        "Valida filas, calcula base = total − IVA y consolida por número/proveedor.",
        "Gastos consolidados e historial de importación.",
        "importaciones y gastos_consolidados (más estructura histórica gastos_raw).")
    add_bullet(doc, "Detecta automáticamente delimitador coma o punto y coma en CSV.")
    add_bullet(doc, "Normaliza encabezados quitando acentos, espacios y símbolos.")
    add_bullet(doc, "Las filas repetidas se consolidan: se suman montos y se amplía el rango de fechas.")
    add_bullet(doc, "El archivo temporal se elimina después del procesamiento.")

    # 15 legacy reconciliation
    doc.add_heading("15. Conciliación factura-vs-gasto — flujo legado", level=1)
    add_callout(doc, "No confundir", "Este módulo compara facturas XML contra gastos consolidados y guarda corridas. Facturas por pagar compara un listado semanal contra facturas XML. Son dos motores relacionados, pero distintos.", ORANGE, GOLD)
    add_module_card(doc,
        "Emparejar facturas XML con gastos consolidados y medir diferencias.",
        "Una importación XML y una importación de gastos, o todos los datos.",
        "Busca candidatos, calcula scores, asigna estados y guarda una corrida auditable.",
        "Cuadrícula de conciliación, revisión manual, ZIP de XML y mapa para PDFs.",
        "conciliacion_corridas, conciliaciones y catalogo_estados.")

    doc.add_heading("15.1 Algoritmo de matching legado", level=2)
    add_table(doc, ["Señal", "Peso", "Cómo se usa"], [
        ("Número", "40%", "Normaliza separadores/ceros, compara núcleo numérico y tolera ciertos formatos largos."),
        ("Proveedor", "30%", "Compara tokens significativos, iniciales y abreviaturas; elimina sufijos societarios."),
        ("Monto", "15%", "Premia totales iguales o cercanos."),
        ("Fecha", "15%", "Premia fechas iguales o próximas."),
    ], [4.0, 2.5, 9.5])
    add_rich_paragraph(doc, ["Primero exige proveedor ≥ 60 para considerar un candidato. Dentro de ese grupo requiere número ≥ 30. Luego calcula el score ponderado y solo acepta el par si el total alcanza 45. Cada gasto puede usarse una sola vez dentro de la corrida."])

    doc.add_heading("15.2 Estados de conciliación", level=2)
    add_table(doc, ["Estado", "Regla resumida"], [
        ("conciliada", "Número 100, proveedor ≥95, fecha ≥85 y diferencia total ≤0.50."),
        ("requiere_revision", "Hay pareja y score total ≥75, pero no cumple todas las condiciones estrictas de conciliada."),
        ("con_diferencias", "Hay pareja, pero la calidad/diferencias no alcanzan el nivel anterior."),
        ("pendiente", "Existe factura XML pero no se encontró gasto razonable."),
        ("gasto_sin_xml", "Existe gasto que no fue utilizado por ninguna factura."),
    ], [4.0, 12.0])
    add_callout(doc, "Verde significa condiciones estrictas", "No basta con que el score general sea alto. El código actual exige simultáneamente número exacto, proveedor casi idéntico, fecha cercana y monto dentro de ₡0.50.", GREEN, "3A8D52")

    doc.add_heading("15.3 Revisión y descargas", level=2)
    add_bullet(doc, "La persona puede cambiar el estado y escribir un comentario; el registro queda marcado como match manual/revisado.")
    add_bullet(doc, "La descarga ZIP incluye XML de facturas con gasto asociado, verifica el hash y evita repetir contenido idéntico.")
    add_bullet(doc, "Los nombres se estandarizan como FE_PROVEEDOR_ddmmaa_numero.")
    add_bullet(doc, "El mapa de nombres para PDF se entrega como JSON; si dos facturas comparten el mismo número corto, se marcan ambiguas y no se renombran automáticamente.")

    # 16 DB
    doc.add_heading("16. Cómo se organiza la información en la base de datos", level=1)
    add_rich_paragraph(doc, ["Una ", ("tabla", {"bold": True}), " se parece a una hoja de cálculo: filas de registros y columnas de propiedades. Una ", ("clave foránea", {"bold": True}), " enlaza una fila con otra tabla mediante su id. El sistema también usa LEFT JOIN para mostrar datos relacionados aunque alguna relación esté vacía."])
    add_table(doc, ["Tabla", "Qué representa", "Relación principal"], [
        ("usuarios", "Personas que ingresan", "La sesión guarda el id del usuario autenticado"),
        ("sociedades", "Empresas del grupo", "Una sola se marca activa"),
        ("semanas", "Períodos de trabajo", "facturas_xml y porpagar_listados tienen semana_id"),
        ("proveedores", "Emisores normalizados", "facturas_xml.proveedor_id"),
        ("facturas_xml", "Comprobantes importados", "Pertenece a proveedor, importación y opcionalmente semana"),
        ("importaciones", "Cabecera/auditoría de una carga", "Agrupa facturas, gastos o items de cola"),
        ("importacion_items", "Archivos individuales de una cola", "Cada item pertenece a una importación"),
        ("porpagar_listados", "Listado de pago de una semana", "Pertenece a sociedad/semana"),
        ("porpagar_facturas", "Cada línea del listado", "Puede apuntar al XML que la respalda"),
        ("correo_cuentas", "Configuraciones IMAP", "Separa índices y bandejas por cuenta"),
        ("correo_indice", "Encabezados buscables", "Cuenta + carpeta + UID identifican un mensaje"),
        ("correo_carpetas", "Progreso de sincronización", "Recuerda último UID/UIDVALIDITY"),
        ("correo_procesados", "Mensajes ya procesados", "Evita descargar/procesar lo mismo"),
        ("correo_bandeja", "XML encontrados esperando decisión", "Puede apuntar a una importación"),
        ("gastos_consolidados", "Gastos agrupados del flujo legado", "Se emparejan con facturas en conciliaciones"),
        ("catalogo_estados", "Estados de conciliación legado", "conciliaciones.estado_id"),
        ("conciliacion_corridas", "Ejecuciones históricas", "Agrupa resultados de conciliaciones"),
        ("conciliaciones", "Resultado factura-vs-gasto", "Puede relacionar ambos lados o solo uno"),
    ], [4.1, 6.2, 5.7])

    doc.add_heading("16.1 Relación principal del flujo actual", level=2)
    add_flow(doc, [
        "sociedades (una activa)",
        "semanas (período elegido)",
        "porpagar_listados → porpagar_facturas (lo que se desea pagar)",
        "facturas_xml → proveedores (lo que respalda el pago)",
        "porpagar_facturas.factura_xml_id une ambos lados",
    ])

    # 17 scenarios
    doc.add_heading("17. Ejemplos completos para entender el recorrido", level=1)
    doc.add_heading("17.1 Ejemplo A: factura localizada en el correo", level=2)
    scenario_a = [
        "Ana inicia sesión y activa la sociedad Empresa A, cédula 3101…",
        "Crea/elige la semana “Pago 20–24 julio”.",
        "Sube el listado por pagar; una línea queda sin respaldo: FACT-71176, Proveedor X, ₡125,000.",
        "Pulsa Buscar en correo. El módulo abre con un término numérico y, si conoce la fecha, prioriza ese mes.",
        "Selecciona el mensaje. Correo descarga los adjuntos y detecta factura XML, MensajeHacienda y PDF.",
        "Comprueba aceptación, receptor, duplicado y empareja el PDF con el XML.",
        "Ana selecciona la fila pendiente e importa a la misma semana.",
        "La cola crea la factura y al terminar re-verifica el listado.",
        "Si número/proveedor/monto cumplen, la línea pasa de sin respaldo a respaldada.",
    ]
    for text in scenario_a:
        add_number(doc, text)

    doc.add_heading("17.2 Ejemplo B: XML cargado manualmente", level=2)
    for text in [
        "Se abre Carga de Facturas XML y se elige la misma semana del listado.",
        "Se arrastran varios XML. La cola crea una importación y procesa tandas.",
        "Un archivo repetido se marca duplicado por hash/consecutivo; los nuevos se guardan.",
        "Al completar la cola, el listado semanal se verifica automáticamente.",
        "Si un XML válido quedó sin usar, aparece en Casos sin coincidencia para moverlo de semana o vincularlo manualmente.",
    ]:
        add_number(doc, text)

    doc.add_heading("17.3 Ejemplo C: diferencia de monto", level=2)
    add_rich_paragraph(doc, ["El listado dice ₡100,000 y el XML dice ₡99,500. Número y proveedor pueden identificar claramente la factura, pero la diferencia es ₡500, mayor que la tolerancia de ₡1. El resultado del módulo por-pagar será ", ("con diferencia", {"bold": True, "color": GOLD}), ". Esto no significa necesariamente que sea el XML equivocado; significa que una persona debe explicar o corregir la diferencia."])

    # 18 mistakes
    doc.add_heading("18. Problemas frecuentes y cómo razonarlos", level=1)
    add_table(doc, ["Síntoma", "Causa probable", "Qué revisar"], [
        ("Todo queda sin respaldo", "Listado y XML están en semanas distintas", "Selector de semana en ambos módulos"),
        ("Correo dice IMAP no disponible", "Extensión PHP IMAP desactivada", "php.ini de XAMPP y reinicio de Apache"),
        ("No aparecen correos recientes", "Índice sin sincronizar, cuenta/carpeta incorrecta o filtro de días", "Cuenta activa, sincronización y rango"),
        ("Factura queda otra_cedula", "Receptor XML distinto a sociedad activa", "Sociedad elegida y NumeroCedulaReceptor"),
        ("Factura rechazada", "MensajeHacienda con código 3", "Detalle de respuesta de Hacienda"),
        ("XML aparece como duplicado", "Mismo hash o mismo comprobante ya registrado", "Historial y factura existente"),
        ("PDF quedó sin identificar", "No hubo XML compatible o el nombre no permitió emparejar", "Correo original y storage/correo/sin_identificar"),
        ("Listado rechaza columnas", "Encabezados faltantes o formato .xls", "Usar Fecha, Numero, Proveedor, Total y guardar como XLSX/CSV"),
        ("Línea con diferencia", "El XML probable existe pero el total difiere > ₡1", "Monto, moneda, notas de crédito y digitación"),
        ("Aviso de variable indefinida en una vista", "El editor analiza la vista sin conocer extract($data)", "Confirmar que el controlador envía la clave; no asumir error de ejecución"),
        ("Error 404", "No existe la combinación de ruta y método GET/POST", "app/config/routes.php"),
        ("Error 500", "Excepción de PHP, base de datos, permisos o archivo", "storage/logs/app.log"),
    ], [4.2, 6.2, 5.6])

    # 19 security/technical
    doc.add_heading("19. Seguridad, integridad y límites importantes", level=1)
    add_bullet(doc, "Consultas preparadas PDO: los valores se envían separados del SQL, reduciendo riesgo de inyección SQL.")
    add_bullet(doc, "htmlspecialchars en vistas: evita que texto de datos se interprete como HTML ejecutable en muchos puntos de salida.")
    add_bullet(doc, "password_hash/password_verify: las contraseñas de usuarios se guardan como hash.")
    add_bullet(doc, "Token CSRF: está implementado explícitamente en el login. No se observa un middleware CSRF general para todos los formularios POST.")
    add_bullet(doc, "Hash de XML: se usa para duplicados e integridad, pero no sustituye una firma digital ni una validación fiscal externa.")
    add_bullet(doc, "Base64 en correo: es codificación reversible, no cifrado.")
    add_bullet(doc, "Creación lazy de tablas: varios modelos intentan CREATE/ALTER al primer uso. En hosting sin permisos DDL se requieren las migraciones SQL manuales.")
    add_bullet(doc, "Entorno: en local se muestran más detalles de error; en producción se ocultan al usuario y se escriben en logs.")
    add_callout(doc, "Principio humano", "Un score alto indica parecido, no verdad. Los estados con diferencia, sin respaldo o revisión deben entenderse como señales para investigar, no como decisiones contables automáticas.", LIGHT_GOLD, GOLD)

    # 20 study map
    doc.add_heading("20. Mapa rápido para estudiar", level=1)
    add_table(doc, ["Si quieres entender…", "Empieza por…", "Archivo técnico principal"], [
        ("Cómo entra una página", "Arquitectura MVC y rutas", "app/config/routes.php + app/core/*"),
        ("Qué empresa se usa", "Inicio y Sociedades", "HomeController.php / Sociedad.php"),
        ("Cómo se separan períodos", "Semanas", "Semana.php + Controller::semanaActiva"),
        ("Cómo se verifica el pago", "Facturas por pagar", "PorPagarController.php + PorPagarVerificador.php"),
        ("Cómo se comparan números/proveedores", "Matching", "FacturaMatcher.php"),
        ("Cómo llegan adjuntos del buzón", "Correo", "CorreoController.php + MailFetcher.php + CorreoSync.php"),
        ("Cómo se interpreta un XML", "Carga XML", "XmlParser.php + InvoiceImportQueue.php"),
        ("Cómo se guardan datos", "Modelos/base de datos", "app/models/* + database/*.sql"),
        ("Cómo se genera Excel", "Reportes", "ReportesController.php + XlsxWriter.php"),
        ("Flujo histórico", "Gastos y Conciliación", "GastosController.php + ConciliacionController.php"),
    ], [5.0, 5.0, 6.0])

    doc.add_heading("20.1 Reglas que conviene memorizar", level=2)
    rules = [
        "La sociedad activa define la cédula receptora esperada.",
        "La semana define el universo del matching por-pagar.",
        "Por-pagar usa solo facturas FE; NC/ND no respaldan una línea de pago.",
        "Número fuerte + proveedor parecido identifican el XML; número muy fuerte + monto exacto puede rescatar nombres distintos.",
        "Hasta ₡1 de diferencia se considera tolerancia en por-pagar.",
        "Un XML duplicado no vuelve a insertarse.",
        "Correo normal busca en un índice local; buscar en todo el cuerpo consulta IMAP y es más lento.",
        "La cola divide cargas grandes en tandas y conserva el estado de cada archivo.",
        "Los vínculos manuales de por-pagar quedan protegidos frente a una re-verificación.",
        "Reportes consulta y exporta; no modifica la lógica del matching.",
    ]
    for item in rules:
        add_bullet(doc, item)

    doc.add_heading("20.2 Preguntas de autoevaluación", level=2)
    questions = [
        "¿Qué diferencia existe entre sociedad activa, semana activa y cuenta de correo activa?",
        "¿Por qué una factura correcta puede quedar sin respaldo si está en otra semana?",
        "¿Qué pasos ocurren entre seleccionar un correo e importar su XML?",
        "¿Qué función cumple el MensajeHacienda y por qué no se importa como otra factura?",
        "¿Qué diferencia hay entre respaldada y con diferencia?",
        "¿En qué situación se usa el rescate por número + monto?",
        "¿Por qué el sistema mantiene una cola de importación?",
        "¿Qué diferencia hay entre Facturas por pagar y Conciliación legado?",
        "¿Qué tabla enlaza una línea por pagar con su factura XML?",
        "¿Qué riesgos no resuelve Base64 en las contraseñas IMAP?",
    ]
    for q in questions:
        add_number(doc, q)

    # 21 source map
    doc.add_heading("21. Mapa técnico de archivos", level=1)
    add_rich_paragraph(doc, ["Esta sección no es necesaria para operar el sistema, pero ayuda a relacionar la explicación con el código."])
    add_table(doc, ["Carpeta/archivo", "Responsabilidad"], [
        ("public/index.php", "Punto de entrada único."),
        ("app/config/routes.php", "Mapa de URLs y acciones."),
        ("app/core/", "Arranque, router, controlador base y modelo base."),
        ("app/controllers/", "Coordinación de cada módulo."),
        ("app/models/", "Acceso a tablas y operaciones de datos."),
        ("app/views/", "Pantallas HTML/PHP."),
        ("app/helpers/XmlParser.php", "Lectura de facturas/MensajeHacienda."),
        ("app/helpers/FacturaMatcher.php", "Reglas compartidas de número/proveedor."),
        ("app/helpers/PorPagarVerificador.php", "Cruce automático listado ↔ XML."),
        ("app/helpers/InvoiceImportQueue.php", "Procesamiento por tandas."),
        ("app/helpers/MailFetcher.php", "Conexión y extracción IMAP."),
        ("app/helpers/CorreoSync.php", "Sincronización incremental del índice."),
        ("database/", "Esquema y migraciones."),
        ("storage/", "Logs y archivos temporales/operativos del correo."),
        ("public/uploads/", "Recepción temporal de archivos subidos."),
    ], [6.2, 9.8])

    doc.add_heading("Cierre", level=1)
    add_rich_paragraph(doc, [
        "La mejor manera de entender Nexo Fiscal es seguir siempre la misma cadena: ",
        ("contexto → entrada → validación → almacenamiento → comparación → estado → acción humana", {"bold": True, "color": NAVY}),
        ". La sociedad y la semana establecen el contexto; Correo/Carga XML aportan evidencia; Facturas por pagar compara; Reportes permite estudiar el resultado."
    ])
    add_callout(doc, "Resumen en una frase", "El sistema organiza comprobantes electrónicos y los cruza con lo que se piensa pagar, automatizando coincidencias claras y dejando visibles los casos que necesitan criterio humano.", GREEN, "3A8D52")

    # Propiedades y guardado
    doc.core_properties.title = "Guía de estudio de Nexo Fiscal"
    doc.core_properties.subject = "Explicación funcional y técnica para principiantes"
    doc.core_properties.author = "Codex, a partir del código del proyecto Nexo Fiscal"
    doc.core_properties.keywords = "Nexo Fiscal, guía, MVC, correo, XML, facturas por pagar, conciliación"
    doc.core_properties.comments = "Generado desde la revisión del código disponible el 20-07-2026."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build_document()
    print(result)
